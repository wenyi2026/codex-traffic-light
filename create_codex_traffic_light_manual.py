from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Codex红绿灯状态提示器-安装与使用说明.docx"
APP_VERSION = "1.2.0"
INSTALLER_NAME = f"CodexTrafficLightInstaller-v{APP_VERSION}.exe"
DARK_PREVIEW = ROOT / "codex_traffic_light_dark_preview.png"
LIGHT_PREVIEW = ROOT / "codex_traffic_light_preview.png"
ICON_PREVIEW = ROOT / "codex_traffic_light_installer_icon.png"


COLORS = {
    "ink": "111827",
    "muted": "5B6472",
    "blue": "2E74B5",
    "blue_dark": "1F4D78",
    "line": "D9E1EA",
    "fill": "F4F7FA",
    "green": "16A34A",
    "yellow": "B7791F",
    "red": "DC2626",
}


def set_east_asian_font(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E1EA", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    set_east_asian_font(run)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix):])
    else:
        add_run(p, text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_east_asian_font(run)
        run.font.color.rgb = RGBColor.from_string(COLORS["blue"] if level <= 2 else COLORS["blue_dark"])
        run.bold = True
    if not p.runs:
        add_run(p, text, bold=True, color=COLORS["blue"] if level <= 2 else COLORS["blue_dark"])
    return p


def add_note_box(doc, title, body, fill="F4F7FA", color="1F4D78"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "D9E1EA", "8")
    p = cell.paragraphs[0]
    add_run(p, title, bold=True, color=color)
    p2 = cell.add_paragraph()
    add_run(p2, body, color=COLORS["ink"])
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, color=COLORS["muted"], size=9)
    return p


def add_key_value_table(doc, rows, widths=(1.55, 4.75)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, widths)
    for row_idx, (key, value) in enumerate(rows):
        c0, c1 = table.rows[row_idx].cells
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c0, "E8EEF5")
        set_cell_border(c0)
        set_cell_border(c1)
        p0 = c0.paragraphs[0]
        add_run(p0, key, bold=True, color=COLORS["blue_dark"])
        p1 = c1.paragraphs[0]
        add_run(p1, value, color=COLORS["ink"])
    return table


def add_status_table(doc):
    headers = ["颜色", "显示状态", "代表含义", "用户应做什么"]
    rows = [
        ["绿色", "GREEN / Idle", "当前没有任务在执行，Codex 已完成任务并处于空闲。", "可以开始新任务，或保持窗口待命。"],
        ["黄色", "YELLOW / Running", "Codex 正在思考、生成回答、运行工具或正常执行任务。", "等待任务完成，通常不需要人工处理。"],
        ["红色", "RED / Needs attention", "需要审批、权限确认、运行异常、任务阻塞或需要人工介入。", "返回 Codex 查看提示，按需批准、修复错误或补充信息。"],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    set_table_width(table, [0.8, 1.35, 2.65, 1.5])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        add_run(cell.paragraphs[0], text, bold=True, color=COLORS["blue_dark"])
    color_map = {"绿色": COLORS["green"], "黄色": COLORS["yellow"], "红色": COLORS["red"]}
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            cell = table.rows[r].cells[c]
            set_cell_border(cell)
            if c == 0:
                add_run(cell.paragraphs[0], text, bold=True, color=color_map[text])
            else:
                add_run(cell.paragraphs[0], text, color=COLORS["ink"])
    return table


def add_image_if_exists(doc, path, width, caption):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)
    else:
        add_note_box(doc, "截图缺失", f"未找到截图文件：{path.name}", fill="FFF7ED", color="B45309")


def apply_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 14, 7),
        ("Heading 3", 12, COLORS["blue_dark"], 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "Codex 红绿灯状态提示器 | 安装与使用说明", color=COLORS["muted"], size=8)


def build():
    doc = Document()
    apply_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "Codex 红绿灯状态提示器", bold=True, color=COLORS["blue_dark"], size=24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "安装与使用说明", color=COLORS["muted"], size=14)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, f"适用系统：Windows | 软件版本：{APP_VERSION} 轻量优化版 | 文档版本：1.2", color=COLORS["muted"], size=9)
    add_image_if_exists(doc, ICON_PREVIEW, 1.8, "安装包图标示意：三色红绿灯图标")

    add_note_box(
        doc,
        "一句话说明",
        "本软件是一个桌面置顶的 Codex 状态提示器：绿色表示空闲，黄色表示正在思考或执行，红色表示需要审批、异常或人工介入。",
    )

    add_heading(doc, "1. 快速安装", 1)
    steps = [
        f"双击安装包 {INSTALLER_NAME}。",
        "如 Windows 弹出安全提醒，请确认来源可信后选择“更多信息 / 仍要运行”。",
        "安装完成后，程序会写入本机用户目录，并创建桌面快捷方式和开机监听项。",
        "之后打开 Codex 时，红绿灯程序会自动关联启动，并始终显示在桌面最上层。",
    ]
    for step in steps:
        add_paragraph(doc, step, style="List Number")

    add_heading(doc, "2. 软件用途", 1)
    add_paragraph(
        doc,
        "它用于让使用者在桌面上直观看到 Codex 当前状态，尤其适合 Codex 窗口被其他软件遮挡、长任务等待、或需要及时处理审批请求的场景。",
    )
    add_key_value_table(
        doc,
        [
            ("核心作用", "用红、黄、绿三种颜色提示 Codex 当前是否空闲、正在运行、或需要人工介入。"),
            ("显示方式", "小型悬浮窗口，默认置顶显示；支持最小化到系统托盘和关闭。"),
            ("主题模式", "支持深色和浅色两种模式；手动切换后会记住上次选择。"),
            ("产品入口", "桌面快捷方式、开始菜单启动项、系统托盘菜单、开始菜单卸载入口。"),
            ("轻量优化", "后台 watcher 只负责等待 Codex 启动；检测到 Codex 后会拉起主界面并退出，减少长期内存占用。"),
            ("自动行为", "Codex 启动后自动打开红绿灯；Codex 完全退出后，红绿灯窗口会自动关闭。"),
        ],
    )

    add_heading(doc, "3. 状态规则", 1)
    add_status_table(doc)
    add_note_box(
        doc,
        "判断原则",
        "黄色覆盖“正在思考”和“正常执行中”；红色覆盖“审批、异常、阻塞、需要人员处理”；只有确认没有任务执行时才显示绿色。",
        fill="FFFBEA",
        color="9A6B00",
    )

    add_heading(doc, "4. 界面与按钮说明", 1)
    add_paragraph(doc, "主界面包含状态标题、当前状态按钮、三条状态卡片，以及右上角的最小化和关闭按钮。")
    add_key_value_table(
        doc,
        [
            ("CODEX 标题", "点击标题区域可在深色模式和浅色模式之间切换；选择会自动保存。"),
            ("状态按钮", "右上方彩色按钮显示当前状态，文字颜色与状态颜色一致。"),
            ("最小化按钮", "点击“-”可把窗口收纳到系统托盘。"),
            ("关闭按钮", "点击“X”可退出红绿灯窗口；下次 Codex 启动时仍可由监听程序自动打开。"),
            ("右下角斜线", "鼠标悬停会变为斜向缩放光标，按住拖动可自由调整窗口大小。"),
            ("右键菜单", "可恢复窗口、切换主题、手动切换状态、调节大小、查看状态、打开关于窗口或退出。"),
            ("系统托盘", "最小化后可从托盘图标右键菜单恢复、切换主题或退出。"),
            ("双击窗口", "可按红、黄、绿顺序循环切换，用于临时检查显示效果。"),
        ],
    )

    add_heading(doc, "5. 功能截图说明", 1)
    add_image_if_exists(doc, DARK_PREVIEW, 6.3, "深色模式：从左到右分别为红色审批/异常、黄色运行/思考、绿色完成/空闲。")
    add_image_if_exists(doc, LIGHT_PREVIEW, 6.3, "浅色模式：布局与深色模式一致，适合浅色桌面或办公环境。")
    add_paragraph(
        doc,
        "截图中的三组面板用于说明三种状态的视觉效果。实际使用时窗口只显示当前状态，不需要用户手动判断日志。",
    )

    add_heading(doc, "6. 自动同步逻辑", 1)
    add_paragraph(doc, "程序会监听 Codex 本地会话事件，并按下列优先级刷新显示：")
    sync_steps = [
        "检测到审批请求、运行异常、阻塞或需要人工介入时，立即显示红色。",
        "检测到 Codex 正在思考、生成回答、调用工具或执行命令时，显示黄色。",
        "检测到任务完成且没有继续执行的活动时，显示绿色。",
        "如果暂时没有读到会话事件，会使用 Codex 进程活动作为辅助判断。进程活动仅作为兜底，不作为最高优先级。",
    ]
    for step in sync_steps:
        add_paragraph(doc, step, style="List Number")
    add_note_box(
        doc,
        "重要说明",
        "手动切换状态主要用于演示或排查显示效果。自动检测开启时，下一次检测到真实 Codex 事件后，界面会重新回到真实状态。",
        fill="F4F7FA",
        color=COLORS["blue_dark"],
    )

    add_heading(doc, "7. 日常使用流程", 1)
    add_heading(doc, "7.1 正常使用", 2)
    normal_steps = [
        "打开 Codex。",
        "观察桌面悬浮红绿灯。",
        "黄色时等待 Codex 执行；红色时返回 Codex 处理审批或异常；绿色时可以开始下一项任务。",
    ]
    for step in normal_steps:
        add_paragraph(doc, step, style="List Number")
    add_heading(doc, "7.2 切换深色 / 浅色模式", 2)
    for step in ["点击窗口左上方 CODEX 标题文字。", "界面会在深色和浅色模式之间切换。", "软件会记住本次选择，下次随 Codex 启动时沿用上次模式。"]:
        add_paragraph(doc, step, style="List Number")
    add_heading(doc, "7.3 关闭程序", 2)
    add_paragraph(doc, "点击右上角 X 可关闭当前窗口。点击“-”会收纳到系统托盘；从托盘右键菜单可恢复或退出。Codex 完全退出后，红绿灯窗口也会自动关闭。")

    add_heading(doc, "8. 常见问题", 1)
    faq_rows = [
        ("安装时提示不安全怎么办？", "这是因为安装包是本地打包的自定义 EXE，未做商业代码签名。确认文件来源可信后，可以选择继续运行。"),
        ("为什么手动切到绿色后又变黄？", "这是正常现象。手动切换只用于临时演示，自动检测到 Codex 正在运行或思考后会恢复真实状态。"),
        ("为什么红色亮起？", "通常表示需要审批、权限确认、命令异常、任务阻塞或需要你返回 Codex 补充信息。"),
        ("关闭后还会不会自动打开？", "如果监听项仍在运行，下次 Codex 启动时会自动打开。"),
        ("最小化后去哪了？", "窗口会收纳到 Windows 系统托盘，可右键托盘图标选择 Show window 恢复。"),
        ("桌面图标没更新怎么办？", "在桌面空白处按 F5 刷新，或重新登录 Windows 后再查看。"),
    ]
    add_key_value_table(doc, faq_rows, widths=(2.0, 4.3))

    add_heading(doc, "9. 文件位置与卸载建议", 1)
    add_key_value_table(
        doc,
        [
            ("安装目录", r"%LOCALAPPDATA%\CodexTrafficLight"),
            ("主程序", r"%LOCALAPPDATA%\CodexTrafficLight\CodexTrafficLight.exe"),
            ("配置文件", r"%LOCALAPPDATA%\CodexTrafficLight\codex_traffic_light.config.json"),
            ("产品信息", r"%LOCALAPPDATA%\CodexTrafficLight\product.json"),
            ("卸载脚本", r"%LOCALAPPDATA%\CodexTrafficLight\Uninstall Codex Traffic Light.cmd"),
            ("自启动项", r"%APPDATA%\Microsoft\Windows\Start Menu\Programs\Startup\Codex Traffic Light Watcher.cmd"),
            ("轻量监听器", r"%LOCALAPPDATA%\CodexTrafficLight\CodexTrafficLightWatcher.exe"),
            ("卸载方式", "从开始菜单 Codex Traffic Light 文件夹中点击 Uninstall Codex Traffic Light，或运行安装目录内的卸载脚本。"),
        ],
        widths=(1.75, 4.55),
    )

    add_heading(doc, "10. 发送给他人时建议附带的内容", 1)
    for item in [
        f"{INSTALLER_NAME} 安装包。",
        "本说明文档。",
        "一句提醒：安装后打开 Codex，红绿灯会自动出现；绿色空闲、黄色运行、红色需处理。",
    ]:
        add_paragraph(doc, item, style="List Bullet")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
